#!/usr/bin/env python3
"""KNOU 소프트웨어경진대회 참가신청서 생성기.

공식 HWP 양식과 같은 표 구조·치수(원본 HWP 레코드 기준, 총폭 17.68cm)로
`turtlemeck-application.docx`와 `turtlemeck-application.pdf`를 만든다.

개인정보는 커밋되지 않는 `participants.env`(KEY=VALUE)에서 읽고,
같은 이름의 환경 변수가 있으면 그 값이 우선한다. 항목은
`participants.env.example`을 참고한다.

사용법:
    python3 -m pip install python-docx
    python3 docs/knou/generate-application.py
PDF 변환에는 Google Chrome이 필요하며, 없으면 docx만 생성한다.
"""

import os
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

HERE = Path(__file__).resolve().parent
ENV_PATH = HERE / "participants.env"
DOCX_OUT = HERE / "turtlemeck-application.docx"
PDF_OUT = HERE / "turtlemeck-application.pdf"
CHROME = "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"

DEFAULTS = {
    "KNOU_REGION": "[소속지역]",
    "KNOU_APPLY_DATE": "2026년   7월   24일",
    "KNOU_LEADER_NAME": "[성명]",
    "KNOU_LEADER_STUDENT_ID": "[학번]",
    "KNOU_LEADER_GRADE": "[학년]",
    "KNOU_LEADER_EMAIL": "[이메일]",
    "KNOU_LEADER_PHONE": "[전화번호]",
    "KNOU_MEMBER_NAME": "[성명]",
    "KNOU_MEMBER_STUDENT_ID": "[학번]",
    "KNOU_MEMBER_GRADE": "[학년]",
    "KNOU_MEMBER_EMAIL": "[이메일]",
    "KNOU_MEMBER_PHONE": "[전화번호]",
}

OVERVIEW = (
    "turtlemeck는 장시간 컴퓨터를 사용하는 사람의 자세 습관 형성을 돕는 macOS 메뉴 막대 앱이다. "
    "별도 장비 없이 Mac 내장 카메라로 짧은 이미지 묶음을 촬영하고, 머리와 몸통의 상대적인 앞뒤 관계를 "
    "기기 안에서 분석한다. 최초 실행 시 바른 자세를 개인 기준으로 저장하고, 이후 결과를 이 기준과 비교해 "
    "좋지 않은 자세가 지속될 때만 알림을 보낸다. 일시적인 움직임이나 판단하기 어려운 입력은 나쁜 자세로 "
    "단정하지 않으며, 운영 모드에서는 촬영 이미지를 저장하거나 외부로 전송하지 않는다. 메뉴 막대에서 현재 "
    "상태와 오늘의 바른 자세·주의 자세 시간, 알림 횟수를 확인할 수 있다.",
    "[개발 배경] 노트북 화면만 들여다보는 생활로 팀원 스스로 목과 어깨 건강이 나빠지는 것을 체감했고, "
    "최근 가족(삼촌)이 목디스크로 수술을 받으면서 바른 자세 습관을 도와주는 도구의 필요성을 절감했다. "
    "가족과 주변 사람들이 모두 MacBook을 사용하는 환경을 고려해, 모든 노트북에 기본으로 있는 웹캠 "
    "하나만으로 동작하는 macOS 전용 앱으로 개발했다.",
    "[수업 내용 응용] 팀원 모두 2026학년도 1학기에 수강한 인공지능 과목에서 학습한 기계학습·신경망의 "
    "기본 개념과 모델 추론·평가 관점을, 2D 자세 추정(PoseNet)과 단안 상대 깊이 추정(Depth Anything V2 "
    "Small) 모델의 선정·결합과 Core ML 온디바이스 실행에 적용했다. 모델 출력을 그대로 신뢰하지 않고 정렬 "
    "기반 순위 통계(중앙값·백분위수·사분위범위)와 상태 기계 설계로 여러 프레임의 대표값 집계, 불안정한 "
    "입력 제외와 정상·주의·판단 보류 상태 전이를 구성했다. 개발 과정에서 인공지능 분야에 관심이 깊어져 "
    "팀장이 2학기 머신러닝·딥러닝 과목을 수강 신청했다.",
)

SOURCES = (
    "- Apple Core ML Depth Anything V2 Small 모델: Apache-2.0 "
    "(https://huggingface.co/apple/coreml-depth-anything-v2-small)",
    "- Apple PoseNet MobileNet(0.75) 모델: Apache-2.0 "
    "(Apple 'Detecting Human Body Poses in an Image' 샘플)",
    "- Apple PoseNet 샘플 코드를 사용한 부분: MIT",
    "- turtlemeck 자체 소스: MIT",
    "※ 앱 번들에 Apache-2.0 전문과 제3자 고지 문서(ThirdPartyNotices.md)를 포함한다.",
)

TITLE = "2026 컴퓨터과학과 총장배 소프트웨어경진대회 참가신청서"

# 원본 HWP 셀 경계를 합친 10열 그리드와 행 높이 (cm)
COLS = [2.16, 1.37, 0.92, 2.71, 1.67, 1.77, 0.45, 3.08, 0.09, 3.47]
ROW_HEIGHTS = [1.91, 1.37, 1.06, 1.06, 1.06, 1.06, 1.06, 1.01, 1.01, 4.84, 1.87, 7.34]
SECTION_ROWS = {1, 7, 9, 10, 11}  # 원본에서 굵은 경계선으로 구분되는 행


def load_values():
    values = dict(DEFAULTS)
    if ENV_PATH.exists():
        for line in ENV_PATH.read_text(encoding="utf-8").splitlines():
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                key, _, value = line.partition("=")
                if key.strip() in DEFAULTS:
                    values[key.strip()] = value.strip()
    for key in DEFAULTS:
        if os.environ.get(key):
            values[key] = os.environ[key]
    return values


def split_student_id(sid):
    # 학번 열이 좁아 하이픈 뒤에서 줄을 나눈다.
    return sid.replace("-", "-\n", 1) if "-" in sid and len(sid) > 10 else sid


def split_region(region, sep):
    return region.replace("(", f"{sep}(", 1) if "(" in region else region


def build_docx(v):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.top_margin = section.bottom_margin = Cm(1.6)
    section.left_margin = section.right_margin = Cm(1.66)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(16)

    table = doc.add_table(rows=12, cols=10)
    table.autofit = False
    for i, w in enumerate(COLS):
        for cell in table.columns[i].cells:
            cell.width = Cm(w)
    grid_el = table._tbl.find(qn("w:tblGrid"))
    for gc, w in zip(grid_el.findall(qn("w:gridCol")), COLS):
        gc.set(qn("w:w"), str(Cm(w).twips))
    for i, h in enumerate(ROW_HEIGHTS):
        table.rows[i].height = Cm(h)
        table.rows[i].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    def table_borders(tbl, outer_sz=16, inner_sz=4):
        tbl_pr = tbl._tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for tag, sz in (("top", outer_sz), ("left", outer_sz), ("bottom", outer_sz),
                        ("right", outer_sz), ("insideH", inner_sz), ("insideV", inner_sz)):
            el = OxmlElement(f"w:{tag}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:color"), "000000")
            borders.append(el)
        tbl_pr.append(borders)

    def top_border(cell, sz=16):
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        el = OxmlElement("w:top")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), "000000")
        borders.append(el)

    table_borders(table)

    merges = [
        ((0, 0), (0, 1)), ((0, 2), (0, 5)), ((0, 6), (0, 7)), ((0, 8), (0, 9)),
        ((1, 0), (1, 9)),
        *[((r, 1), (r, 2)) for r in range(2, 7)],
        *[((r, 5), (r, 8)) for r in range(2, 7)],
        ((7, 0), (8, 0)), ((7, 1), (8, 4)), ((7, 5), (7, 6)), ((7, 7), (7, 9)),
        ((8, 5), (8, 6)), ((8, 7), (8, 9)),
        ((9, 1), (9, 9)),
        ((10, 0), (10, 2)), ((10, 3), (10, 9)),
        ((11, 0), (11, 9)),
    ]
    for (r0, c0), (r1, c1) in merges:
        table.cell(r0, c0).merge(table.cell(r1, c1))

    for r in SECTION_ROWS:
        seen = set()
        for cell in table.rows[r].cells:
            if id(cell._tc) not in seen:
                seen.add(id(cell._tc))
                top_border(cell)

    def put(r, c, text, *, bold=False, size=10):
        cell = table.cell(r, c)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for i, line in enumerate(text.split("\n") if text else [""]):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.bold = bold
            run.font.size = Pt(size)
        return cell

    put(0, 0, "작품명", bold=True)
    put(0, 2, "turtlemeck")
    put(0, 6, "소속지역", bold=True)
    put(0, 8, split_region(v["KNOU_REGION"], "\n"), size=9)

    put(1, 0, "인 적 사 항", bold=True)
    put(2, 1, "학 번", bold=True)
    put(2, 3, "학 년", bold=True)
    put(2, 4, "성 명", bold=True)
    put(2, 5, "e-mail\n(※수신가능이메일주소)", bold=True, size=8.5)
    put(2, 9, "전화번호\n(※연락가능번호)", bold=True, size=8.5)

    put(3, 0, "팀원(발표자)", size=9)
    put(3, 1, split_student_id(v["KNOU_LEADER_STUDENT_ID"]), size=9)
    put(3, 3, v["KNOU_LEADER_GRADE"], size=9)
    put(3, 4, v["KNOU_LEADER_NAME"], size=9)
    put(3, 5, v["KNOU_LEADER_EMAIL"], size=9)
    put(3, 9, v["KNOU_LEADER_PHONE"], size=9)

    put(4, 0, "팀원1", size=9)
    put(4, 1, split_student_id(v["KNOU_MEMBER_STUDENT_ID"]), size=9)
    put(4, 3, v["KNOU_MEMBER_GRADE"], size=9)
    put(4, 4, v["KNOU_MEMBER_NAME"], size=9)
    put(4, 5, v["KNOU_MEMBER_EMAIL"], size=9)
    put(4, 9, v["KNOU_MEMBER_PHONE"], size=9)

    put(5, 0, "팀원2", size=9)
    put(6, 0, "팀원3", size=9)

    put(7, 0, "출품분야", bold=True)
    put(7, 1, "유틸리티 앱")
    put(7, 5, "운영체제", bold=True)
    put(7, 7, "macOS 15.0 이상")
    put(8, 5, "사용언어", bold=True)
    put(8, 7, "Swift")

    put(9, 0, "작품개요", bold=True)
    cell = table.cell(9, 1)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, para in enumerate(OVERVIEW):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(para)
        run.font.size = Pt(9.5)
        if i:
            p.paragraph_format.space_before = Pt(6)

    put(10, 0, "출 처\n(오픈소스 사용내역,\n라이선스 정보 등)", bold=True, size=8.5)
    cell = table.cell(10, 3)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, line in enumerate(SOURCES):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(line).font.size = Pt(9)

    # 날짜·서명: 날짜 중앙 → 신청인 오른쪽 → 학과 하단 중앙 (원본 배치)
    cell = table.cell(11, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(42)
    p.add_run(v["KNOU_APPLY_DATE"]).font.size = Pt(12)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(46)
    p.paragraph_format.right_indent = Cm(1.2)
    p.add_run(f"참가 신청인   {v['KNOU_LEADER_NAME']}          (인)").font.size = Pt(12)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    run = p.add_run("컴  퓨  터  과  학  과")
    run.font.size = Pt(13)
    run.bold = True

    doc.save(DOCX_OUT)
    print("saved:", DOCX_OUT)


HTML_HEAD = """<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="utf-8">
<style>
  @page { size: A4; margin: 16mm 16.6mm; }
  body { font-family: 'Apple SD Gothic Neo', sans-serif; color: #000; margin: 0; }
  h1 { font-size: 16pt; text-align: center; margin: 8px 0 16px; letter-spacing: -0.3px; }
  table { border-collapse: collapse; width: 100%; table-layout: fixed; font-size: 10pt; border: 1.6pt solid #000; }
  td { border: 0.6pt solid #000; padding: 3px 6px; text-align: center; vertical-align: middle; word-break: keep-all; }
  tr { page-break-inside: avoid; }
  tr.sec > td { border-top: 1.6pt solid #000; }
  td.left { text-align: left; }
  td.head { font-weight: 700; }
  .small { font-size: 8.5pt; }
  .data { font-size: 9pt; }
  .overview p { margin: 0 0 7px; text-align: left; font-size: 9.5pt; line-height: 1.55; }
  .overview p:last-child { margin-bottom: 0; }
  .sources p { margin: 0 0 3px; text-align: left; font-size: 9pt; line-height: 1.5; }
  .sign { vertical-align: top; padding: 0; }
  .sign .date { text-align: center; font-size: 12pt; margin-top: 1.5cm; }
  .sign .applicant { text-align: right; font-size: 12pt; margin-top: 1.6cm; padding-right: 1.2cm; }
  .sign .dept { text-align: center; font-size: 13pt; font-weight: 700; letter-spacing: 6px; margin-top: 1.4cm; }
</style>
</head>
<body>
"""


def build_pdf(v):
    if not Path(CHROME).exists():
        print("Google Chrome이 없어 PDF 생성을 건너뜀. docx를 Word에서 열어 PDF로 내보낼 것.")
        return

    def para(text):
        return "".join(f"      <p>{line}</p>\n" for line in text)

    member_row = (
        '  <tr style="height:1.06cm">\n'
        '    <td class="data">{label}</td>\n'
        '    <td colspan="2" class="data">{sid}</td>\n'
        '    <td class="data">{grade}</td>\n'
        '    <td class="data">{name}</td>\n'
        '    <td colspan="4" class="data">{email}</td>\n'
        '    <td class="data">{phone}</td>\n'
        "  </tr>\n"
    )
    empty_row = (
        '  <tr style="height:1.06cm">\n'
        '    <td class="data">{label}</td><td colspan="2"></td><td></td><td></td>'
        "<td colspan=\"4\"></td><td></td>\n  </tr>\n"
    )

    html = HTML_HEAD
    html += f"<h1>{TITLE}</h1>\n<table>\n"
    html += (
        "  <colgroup>\n"
        '    <col style="width:12.22%"><col style="width:7.74%"><col style="width:5.18%">'
        '<col style="width:15.33%"><col style="width:9.42%">\n'
        '    <col style="width:9.99%"><col style="width:2.54%"><col style="width:17.41%">'
        '<col style="width:0.52%"><col style="width:19.65%">\n'
        "  </colgroup>\n"
    )
    html += (
        '  <tr style="height:1.91cm">\n'
        '    <td colspan="2" class="head">작품명</td>\n'
        '    <td colspan="4">turtlemeck</td>\n'
        '    <td colspan="2" class="head">소속지역</td>\n'
        f'    <td colspan="2" class="data">{split_region(v["KNOU_REGION"], "<br>")}</td>\n'
        "  </tr>\n"
        '  <tr class="sec" style="height:1.37cm"><td colspan="10" class="head">인 적 사 항</td></tr>\n'
        '  <tr style="height:1.06cm">\n'
        "    <td></td>\n"
        '    <td colspan="2" class="head">학 번</td>\n'
        '    <td class="head">학 년</td>\n'
        '    <td class="head">성 명</td>\n'
        '    <td colspan="4" class="head">e-mail<br><span class="small">(※수신가능이메일주소)</span></td>\n'
        '    <td class="head">전화번호<br><span class="small">(※연락가능번호)</span></td>\n'
        "  </tr>\n"
    )
    html += member_row.format(
        label="팀원(발표자)", sid=split_student_id(v["KNOU_LEADER_STUDENT_ID"]).replace("\n", "<br>"),
        grade=v["KNOU_LEADER_GRADE"], name=v["KNOU_LEADER_NAME"],
        email=v["KNOU_LEADER_EMAIL"], phone=v["KNOU_LEADER_PHONE"])
    html += member_row.format(
        label="팀원1", sid=split_student_id(v["KNOU_MEMBER_STUDENT_ID"]).replace("\n", "<br>"),
        grade=v["KNOU_MEMBER_GRADE"], name=v["KNOU_MEMBER_NAME"],
        email=v["KNOU_MEMBER_EMAIL"], phone=v["KNOU_MEMBER_PHONE"])
    html += empty_row.format(label="팀원2")
    html += empty_row.format(label="팀원3")
    html += (
        '  <tr class="sec" style="height:1.01cm">\n'
        '    <td rowspan="2" class="head">출품분야</td>\n'
        '    <td rowspan="2" colspan="4">유틸리티 앱</td>\n'
        '    <td colspan="2" class="head">운영체제</td>\n'
        '    <td colspan="3">macOS 15.0 이상</td>\n'
        "  </tr>\n"
        '  <tr style="height:1.01cm">\n'
        '    <td colspan="2" class="head">사용언어</td>\n'
        '    <td colspan="3">Swift</td>\n'
        "  </tr>\n"
        '  <tr class="sec">\n'
        '    <td class="head">작품개요</td>\n'
        '    <td colspan="9" class="left overview" style="padding:7px 9px">\n'
        f"{para(OVERVIEW)}"
        "    </td>\n"
        "  </tr>\n"
        '  <tr class="sec" style="height:1.87cm">\n'
        '    <td colspan="3" class="head">출 처<br><span class="small">(오픈소스 사용내역,<br>라이선스 정보 등)</span></td>\n'
        '    <td colspan="7" class="left sources" style="padding:6px 9px">\n'
        f"{para(SOURCES)}"
        "    </td>\n"
        "  </tr>\n"
        '  <tr class="sec" style="height:7.34cm">\n'
        '    <td colspan="10" class="sign">\n'
        f'      <div class="date">{v["KNOU_APPLY_DATE"].replace("   ", "&nbsp;&nbsp;&nbsp;")}</div>\n'
        f'      <div class="applicant">참가 신청인&nbsp;&nbsp;&nbsp;{v["KNOU_LEADER_NAME"]}'
        "&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;(인)</div>\n"
        '      <div class="dept">컴 퓨 터 과 학 과</div>\n'
        "    </td>\n"
        "  </tr>\n</table>\n</body>\n</html>\n"
    )

    with tempfile.NamedTemporaryFile("w", suffix=".html", encoding="utf-8", delete=False) as f:
        f.write(html)
        html_path = f.name
    try:
        subprocess.run(
            [CHROME, "--headless", "--disable-gpu", "--no-pdf-header-footer",
             f"--print-to-pdf={PDF_OUT}", html_path],
            check=True, capture_output=True)
        print("saved:", PDF_OUT)
    finally:
        os.unlink(html_path)


def main():
    values = load_values()
    missing = [k for k, val in values.items() if val == DEFAULTS[k] and val.startswith("[")]
    if missing:
        print("주의: 다음 항목이 placeholder 상태다 —", ", ".join(missing))
        print(f"{ENV_PATH.name}를 작성하거나 환경 변수를 설정할 것 (participants.env.example 참고).")
    build_docx(values)
    build_pdf(values)


if __name__ == "__main__":
    main()
